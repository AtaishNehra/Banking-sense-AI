"""
Document processing service for PDF/DOCX summarization and compliance extraction.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict
import pdfplumber
import docx
from transformers import pipeline
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Service for processing and summarizing banking documents.
    """
    
    def __init__(self):
        self.summarizer = None
        self.load_summarization_model()
    
    def load_summarization_model(self):
        """
        Load summarization model from Hugging Face.
        """
        try:
            # Try to load a good summarization model
            model_options = [
                "facebook/bart-large-cnn",
                "t5-small",
                "sshleifer/distilbart-cnn-12-6"
            ]
            
            for model_name in model_options:
                try:
                    logger.info(f"Loading summarization model: {model_name}")
                    self.summarizer = pipeline(
                        "summarization", 
                        model=model_name,
                        device=-1  # Use CPU
                    )
                    logger.info(f"Successfully loaded model: {model_name}")
                    break
                except Exception as e:
                    logger.warning(f"Failed to load {model_name}: {e}")
                    continue
            
            if self.summarizer is None:
                logger.warning("Failed to load any summarization model, using fallback")
                
        except Exception as e:
            logger.error(f"Error loading summarization model: {e}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using pdfplumber.
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from supported file formats.
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def summarize_text(self, text: str, max_length: int = 150, min_length: int = 50) -> str:
        """
        Summarize text using the loaded model.
        """
        if not text.strip():
            return "No text content found in document."
        
        try:
            if self.summarizer:
                # Split text into chunks if too long
                max_chunk_length = 1024  # Model input limit
                chunks = []
                
                if len(text) > max_chunk_length:
                    # Split into sentences and group into chunks
                    sentences = text.split('. ')
                    current_chunk = ""
                    
                    for sentence in sentences:
                        if len(current_chunk + sentence) < max_chunk_length:
                            current_chunk += sentence + ". "
                        else:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            current_chunk = sentence + ". "
                    
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                else:
                    chunks = [text]
                
                # Summarize each chunk
                summaries = []
                for chunk in chunks:
                    if len(chunk.strip()) > 10:  # Skip very short chunks
                        try:
                            summary = self.summarizer(
                                chunk,
                                max_length=max_length,
                                min_length=min(min_length, len(chunk.split()) // 2),
                                do_sample=False
                            )
                            summaries.append(summary[0]['summary_text'])
                        except Exception as e:
                            logger.warning(f"Failed to summarize chunk: {e}")
                            summaries.append(chunk[:200] + "...")
                
                # Combine summaries
                if len(summaries) > 1:
                    combined_summary = " ".join(summaries)
                    # If combined summary is still too long, summarize again
                    if len(combined_summary) > max_length * 2:
                        try:
                            final_summary = self.summarizer(
                                combined_summary,
                                max_length=max_length,
                                min_length=min_length,
                                do_sample=False
                            )
                            return final_summary[0]['summary_text']
                        except:
                            return combined_summary[:max_length * 2] + "..."
                    return combined_summary
                else:
                    return summaries[0] if summaries else "Unable to generate summary."
            
            else:
                # Fallback: simple extractive summarization
                sentences = text.split('. ')
                if len(sentences) <= 3:
                    return text
                
                # Take first and last sentences, and one from middle
                summary_sentences = [
                    sentences[0],
                    sentences[len(sentences) // 2],
                    sentences[-1]
                ]
                return ". ".join(summary_sentences) + "."
                
        except Exception as e:
            logger.error(f"Error in text summarization: {e}")
            # Return first few sentences as fallback
            sentences = text.split('. ')[:3]
            return ". ".join(sentences) + "."
    
    def extract_aml_covenants(self, text: str) -> List[str]:
        """
        Extract AML (Anti-Money Laundering) covenants from text.
        """
        # Keywords and phrases related to AML compliance
        aml_keywords = [
            "anti-money laundering", "aml", "money laundering",
            "suspicious activity", "customer due diligence", "cdd",
            "know your customer", "kyc", "beneficial ownership",
            "politically exposed person", "pep", "sanctions",
            "suspicious transaction", "currency transaction report",
            "ctr", "suspicious activity report", "sar",
            "compliance program", "risk assessment", "monitoring",
            "record keeping", "reporting requirements"
        ]
        
        try:
            # Split text into sentences
            sentences = text.split('.')
            aml_clauses = []
            
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) < 20:  # Skip very short sentences
                    continue
                
                # Check if sentence contains AML-related keywords
                sentence_lower = sentence.lower()
                if any(keyword in sentence_lower for keyword in aml_keywords):
                    aml_clauses.append(sentence)
            
            # Remove duplicates while preserving order
            unique_clauses = []
            for clause in aml_clauses:
                if clause not in unique_clauses:
                    unique_clauses.append(clause)
            
            logger.info(f"Found {len(unique_clauses)} AML-related clauses")
            return unique_clauses
            
        except Exception as e:
            logger.error(f"Error extracting AML covenants: {e}")
            return []
    
    def process_document(self, file_path: str) -> Dict:
        """
        Complete document processing pipeline.
        """
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            if not text.strip():
                return {
                    "summary": "No readable text found in document.",
                    "key_clauses": [],
                    "error": "Empty document"
                }
            
            # Summarize text
            summary = self.summarize_text(text)
            
            # Extract AML covenants
            key_clauses = self.extract_aml_covenants(text)
            
            # Save outputs
            self.save_outputs(file_path, summary, key_clauses, text)
            
            return {
                "summary": summary,
                "key_clauses": key_clauses
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                "summary": f"Error processing document: {str(e)}",
                "key_clauses": [],
                "error": str(e)
            }
    
    def save_outputs(self, file_path: str, summary: str, key_clauses: List[str], full_text: str):
        """
        Save processing outputs to files.
        """
        try:
            # Create output directories
            summaries_dir = Path("data/processed/summaries")
            samples_dir = Path("docs/sample_outputs")
            summaries_dir.mkdir(parents=True, exist_ok=True)
            samples_dir.mkdir(parents=True, exist_ok=True)
            
            # Get base filename
            base_name = Path(file_path).stem
            
            # Save summary
            summary_path = summaries_dir / f"{base_name}_summary.txt"
            with open(summary_path, 'w', encoding='utf-8') as f:
                f.write(summary)
            
            # Save key clauses
            clauses_path = summaries_dir / f"{base_name}_aml_clauses.json"
            with open(clauses_path, 'w', encoding='utf-8') as f:
                json.dump(key_clauses, f, indent=2)
            
            # Save sample output
            sample_path = samples_dir / f"{base_name}_processed.json"
            sample_output = {
                "document": base_name,
                "summary": summary,
                "aml_clauses_count": len(key_clauses),
                "sample_clauses": key_clauses[:3],  # First 3 clauses as sample
                "text_length": len(full_text)
            }
            
            with open(sample_path, 'w', encoding='utf-8') as f:
                json.dump(sample_output, f, indent=2)
            
            logger.info(f"Saved outputs for {base_name}")
            
        except Exception as e:
            logger.error(f"Error saving outputs: {e}")
